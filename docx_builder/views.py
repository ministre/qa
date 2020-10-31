from django.urls import reverse
from django.views.generic import CreateView, ListView, UpdateView, DeleteView
from django.utils.decorators import method_decorator
from django.contrib.auth.decorators import login_required
from .models import DocxProfile
from .forms import DocxProfileForm
from django.shortcuts import get_object_or_404
from feature.models import FeatureList, FeatureListCategory, FeatureListItem
from testplan.models import Testplan, Chapter, Category, Test, TestLink, TestChecklist, TestChecklistItem, TestConfig, \
    TestImage, TestComment
from protocol.models import Protocol, ProtocolDevice, ProtocolTestResult, TestResultIssue, TestResultComment
from device.models import Device
import os
from django.conf import settings
from django.http import HttpResponse, Http404
from django.utils import timezone
from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from django.utils.datastructures import MultiValueDictKeyError


@method_decorator(login_required, name='dispatch')
class DocxProfileCreate(CreateView):
    model = DocxProfile
    form_class = DocxProfileForm
    template_name = 'docx_builder/create.html'

    def get_initial(self):
        return {'created_by': self.request.user, 'updated_by': self.request.user}

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['back_url'] = reverse('docx_profiles')
        return context

    def get_success_url(self):
        return reverse('docx_profiles')


@method_decorator(login_required, name='dispatch')
class DocxProfileUpdate(UpdateView):
    model = DocxProfile
    form_class = DocxProfileForm
    template_name = 'docx_builder/update.html'

    def get_initial(self):
        return {'updated_by': self.request.user, 'updated_at': timezone.now()}

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['back_url'] = reverse('docx_profiles')
        return context

    def get_success_url(self):
        self.object.update_timestamp(user=self.request.user)
        return reverse('docx_profiles')


@method_decorator(login_required, name='dispatch')
class DocxProfileDelete(DeleteView):
    model = DocxProfile
    template_name = 'docx_builder/delete.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['back_url'] = reverse('docx_profiles')
        return context

    def get_success_url(self):
        return reverse('docx_profiles')


@method_decorator(login_required, name='dispatch')
class DocxProfileListView(ListView):
    context_object_name = 'docx_profiles'
    queryset = DocxProfile.objects.all()
    template_name = 'docx_builder/list.html'


def shade_cells(cells, shade):
    for cell in cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcVAlign = OxmlElement("w:shd")
        tcVAlign.set(qn("w:fill"), shade)
        tcPr.append(tcVAlign)


def build_document(profile: DocxProfile):
    document = Document()
    document.styles['Title'].font.name = profile.title_font_name
    document.styles['Title'].font.color.rgb = RGBColor(profile.title_font_color_red, profile.title_font_color_green,
                                                       profile.title_font_color_blue)
    document.styles['Title'].font.size = Pt(profile.title_font_size)
    document.styles['Title'].font.bold = profile.title_font_bold
    document.styles['Title'].font.italic = profile.title_font_italic
    document.styles['Title'].font.underline = profile.title_font_underline
    document.styles['Title'].paragraph_format.space_before = Pt(profile.title_space_before)
    document.styles['Title'].paragraph_format.space_after = Pt(profile.title_space_after)
    if profile.title_alignment == 0:
        document.styles['Title'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif profile.title_alignment == 1:
        document.styles['Title'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif profile.title_alignment == 2:
        document.styles['Title'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif profile.title_alignment == 3:
        document.styles['Title'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    document.styles['Heading 1'].font.name = profile.h1_font_name
    document.styles['Heading 1'].font.color.rgb = RGBColor(profile.h1_font_color_red, profile.h1_font_color_green,
                                                           profile.h1_font_color_blue)
    document.styles['Heading 1'].font.size = Pt(profile.h1_font_size)
    document.styles['Heading 1'].font.bold = profile.h1_font_bold
    document.styles['Heading 1'].font.italic = profile.h1_font_italic
    document.styles['Heading 1'].font.underline = profile.h1_font_underline
    document.styles['Heading 1'].paragraph_format.space_before = Pt(profile.h1_space_before)
    document.styles['Heading 1'].paragraph_format.space_after = Pt(profile.h1_space_after)
    if profile.h1_alignment == 0:
        document.styles['Heading 1'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif profile.h1_alignment == 1:
        document.styles['Heading 1'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif profile.h1_alignment == 2:
        document.styles['Heading 1'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif profile.h1_alignment == 3:
        document.styles['Heading 1'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    document.styles['Heading 2'].font.name = profile.h2_font_name
    document.styles['Heading 2'].font.color.rgb = RGBColor(profile.h2_font_color_red, profile.h2_font_color_green,
                                                           profile.h2_font_color_blue)
    document.styles['Heading 2'].font.size = Pt(profile.h2_font_size)
    document.styles['Heading 2'].font.bold = profile.h2_font_bold
    document.styles['Heading 2'].font.italic = profile.h2_font_italic
    document.styles['Heading 2'].font.underline = profile.h2_font_underline
    document.styles['Heading 2'].paragraph_format.space_before = Pt(profile.h2_space_before)
    document.styles['Heading 2'].paragraph_format.space_after = Pt(profile.h2_space_after)
    if profile.h2_alignment == 0:
        document.styles['Heading 2'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif profile.h2_alignment == 1:
        document.styles['Heading 2'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif profile.h2_alignment == 2:
        document.styles['Heading 2'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif profile.h2_alignment == 3:
        document.styles['Heading 2'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    document.styles['Heading 3'].font.name = profile.h3_font_name
    document.styles['Heading 3'].font.color.rgb = RGBColor(profile.h3_font_color_red, profile.h3_font_color_green,
                                                           profile.h3_font_color_blue)
    document.styles['Heading 3'].font.size = Pt(profile.h3_font_size)
    document.styles['Heading 3'].font.bold = profile.h3_font_bold
    document.styles['Heading 3'].font.italic = profile.h3_font_italic
    document.styles['Heading 3'].font.underline = profile.h3_font_underline
    document.styles['Heading 3'].paragraph_format.space_before = Pt(profile.h3_space_before)
    document.styles['Heading 3'].paragraph_format.space_after = Pt(profile.h3_space_after)
    if profile.h3_alignment == 0:
        document.styles['Heading 3'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif profile.h3_alignment == 1:
        document.styles['Heading 3'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif profile.h3_alignment == 2:
        document.styles['Heading 3'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif profile.h3_alignment == 3:
        document.styles['Heading 3'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    document.styles['Normal'].font.name = profile.normal_font_name
    document.styles['Normal'].font.color.rgb = RGBColor(profile.normal_font_color_red, profile.normal_font_color_green,
                                                        profile.normal_font_color_blue)
    document.styles['Normal'].font.size = Pt(profile.normal_font_size)
    document.styles['Normal'].font.bold = profile.normal_font_bold
    document.styles['Normal'].font.italic = profile.normal_font_italic
    document.styles['Normal'].font.underline = profile.normal_font_underline
    document.styles['Normal'].paragraph_format.space_before = Pt(profile.normal_space_before)
    document.styles['Normal'].paragraph_format.space_after = Pt(profile.normal_space_after)
    if profile.normal_alignment == 0:
        document.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif profile.normal_alignment == 1:
        document.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif profile.normal_alignment == 2:
        document.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif profile.normal_alignment == 3:
        document.styles['Normal'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    document.styles['Caption'].font.name = profile.caption_font_name
    document.styles['Caption'].font.color.rgb = RGBColor(profile.caption_font_color_red,
                                                         profile.caption_font_color_green,
                                                         profile.caption_font_color_blue)
    document.styles['Caption'].font.size = Pt(profile.caption_font_size)
    document.styles['Caption'].font.bold = profile.caption_font_bold
    document.styles['Caption'].font.italic = profile.caption_font_italic
    document.styles['Caption'].font.underline = profile.caption_font_underline
    document.styles['Caption'].paragraph_format.space_before = Pt(profile.caption_space_before)
    document.styles['Caption'].paragraph_format.space_after = Pt(profile.caption_space_after)
    if profile.caption_alignment == 0:
        document.styles['Caption'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif profile.caption_alignment == 1:
        document.styles['Caption'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif profile.caption_alignment == 2:
        document.styles['Caption'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif profile.caption_alignment == 3:
        document.styles['Caption'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    document.styles['Quote'].font.name = profile.quote_font_name
    document.styles['Quote'].font.color.rgb = RGBColor(profile.quote_font_color_red,
                                                       profile.quote_font_color_green,
                                                       profile.quote_font_color_blue)
    document.styles['Quote'].font.size = Pt(profile.quote_font_size)
    document.styles['Quote'].font.bold = profile.quote_font_bold
    document.styles['Quote'].font.italic = profile.quote_font_italic
    document.styles['Quote'].font.underline = profile.quote_font_underline
    document.styles['Quote'].paragraph_format.space_before = Pt(profile.quote_space_before)
    document.styles['Quote'].paragraph_format.space_after = Pt(profile.quote_space_after)
    if profile.quote_alignment == 0:
        document.styles['Quote'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif profile.quote_alignment == 1:
        document.styles['Quote'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif profile.quote_alignment == 2:
        document.styles['Quote'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif profile.quote_alignment == 3:
        document.styles['Quote'].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    return document


@login_required
def build_feature_list(request):
    if request.method == 'POST':
        fl = get_object_or_404(FeatureList, id=request.POST['feature_list'])
        document = build_document(profile=DocxProfile.objects.get(id=request.POST['profile']))
        # title
        document.add_paragraph(fl.name, style='Title')
        # items
        categories = FeatureListCategory.objects.filter(feature_list=fl).order_by('id')
        for i, category in enumerate(categories):
            document.add_heading(str(i+1) + '. ' + category.name, level=1)
            items = FeatureListItem.objects.filter(category=category).order_by('id')
            for j, item in enumerate(items):
                document.add_heading(str(i+1) + '.' + str(j+1) + '. ' + item.name, level=2)

        fl_filename = settings.MEDIA_ROOT + '/fl_' + str(fl.id) + '.docx'
        document.save(fl_filename)
        file_path = os.path.join(settings.MEDIA_ROOT, fl_filename)
        if os.path.exists(file_path):
            with open(file_path, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type="application/vnd.ms-word")
                response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_path)
                return response
        raise Http404


@login_required
def build_testplan(request):
    if request.method == 'POST':
        testplan = get_object_or_404(Testplan, id=request.POST['testplan'])
        profile = DocxProfile.objects.get(id=request.POST['profile'])
        document = build_document(profile=profile)
        # title
        document.add_paragraph(testplan.name, style='Title')
        # page header
        try:
            if request.POST['page_header']:
                style = document.styles.add_style('Header Table', WD_STYLE_TYPE.TABLE)
                style.base_style = document.styles['Table Grid']
                style.font.name = 'Cambria'
                style.font.size = Pt(11)
                style.paragraph_format.space_before = Pt(2)
                style.paragraph_format.space_after = Pt(2)
                style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                section = document.sections
                section[0].left_margin = Cm(2.5)
                section[0].right_margin = Cm(1)
                section[0].top_margin = Cm(4.5)
                section[0].bottom_margin = Cm(1)
                header = document.sections[0].header
                table = header.add_table(rows=0, cols=3, width=Cm(19.5))
                table.style = 'Header Table'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                row_cells = table.add_row().cells
                paragraph = row_cells[0].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(profile.header_logo, width=Inches(1.25))
                paragraph = row_cells[1].paragraphs[0]
                run = paragraph.add_run()
                run.add_text(testplan.name)
                row_cells = table.add_row().cells
                paragraph = row_cells[0].paragraphs[0]
                run = paragraph.add_run()
                run.add_text('Редакция: ' + testplan.version)
                paragraph = row_cells[1].paragraphs[0]
                run = paragraph.add_run()
                run.add_text(profile.header_subtitle)
                a = table.cell(0, 1)
                b = table.cell(0, 2)
                a.merge(b)
                table.cell(0, 0).width = Cm(5)
                table.cell(1, 0).width = Cm(5)
                table.cell(0, 1).width = Cm(10.5)
                table.cell(1, 1).width = Cm(10.5)
                table.cell(0, 2).width = Cm(4)
                table.cell(1, 2).width = Cm(4)
                table.cell(0, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        except MultiValueDictKeyError:
            pass

        try:
            if request.POST['convert_textile']:
                convert_textile = True
            else:
                convert_textile = False
        except MultiValueDictKeyError:
            convert_textile = False

        # chapters
        try:
            if request.POST['chapters']:
                chapters = Chapter.objects.filter(testplan=testplan).order_by('id')
                if chapters:
                    for chapter in chapters:
                        document.add_heading(chapter.name, level=1)

                        if convert_textile:
                            paragraphs = chapter.text.split('\r\n')
                            for paragraph in paragraphs:
                                if paragraph.startswith('#'):
                                    if paragraph.startswith('###'):
                                        p = document.add_paragraph(paragraph[4:], style='List Number 3')
                                    elif paragraph.startswith('##'):
                                        p = document.add_paragraph(paragraph[3:], style='List Number 2')
                                    else:
                                        p = document.add_paragraph(paragraph[2:], style='List Number')
                                elif paragraph.startswith('*'):
                                    if paragraph.startswith('***'):
                                        p = document.add_paragraph(paragraph[4:], style='List Bullet 3')
                                    elif paragraph.startswith('**'):
                                        p = document.add_paragraph(paragraph[3:], style='List Bullet 2')
                                    else:
                                        p = document.add_paragraph(paragraph[2:], style='List Bullet')
                                elif paragraph.startswith('> *%{color:'):
                                    p = document.add_paragraph(paragraph[paragraph.find('}')+1:-2], style='Quote')
                                else:
                                    p = document.add_paragraph(paragraph, style='Normal')
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        else:
                            p = document.add_paragraph(chapter.text, style='Normal')
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        except MultiValueDictKeyError:
            pass

        # tests
        categories = Category.objects.filter(testplan=testplan).order_by('id')
        for i, category in enumerate(categories):
            document.add_heading(str(i+1) + '. ' + category.name, level=1)

            tests = Test.objects.filter(category=category).order_by('id')
            for j, test in enumerate(tests):
                document.add_heading(str(i+1) + '.' + str(j+1) + '. ' + test.name, level=2)

                # purpose
                try:
                    if request.POST['purpose']:
                        document.add_heading('Цель', level=3)
                        paragraph = document.add_paragraph(test.purpose, style='Normal')
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                except MultiValueDictKeyError:
                    pass

                # procedure
                try:
                    if request.POST['procedure']:
                        document.add_heading('Процедура', level=3)

                        if convert_textile:
                            paragraphs = test.procedure.split('\r\n')
                            for paragraph in paragraphs:
                                if paragraph.startswith('#'):
                                    if paragraph.startswith('###'):
                                        p = document.add_paragraph(paragraph[4:], style='List Number 3')
                                    elif paragraph.startswith('##'):
                                        p = document.add_paragraph(paragraph[3:], style='List Number 2')
                                    else:
                                        p = document.add_paragraph(paragraph[2:], style='List Number')
                                elif paragraph.startswith('*'):
                                    if paragraph.startswith('***'):
                                        p = document.add_paragraph(paragraph[4:], style='List Bullet 3')
                                    elif paragraph.startswith('**'):
                                        p = document.add_paragraph(paragraph[3:], style='List Bullet 2')
                                    else:
                                        p = document.add_paragraph(paragraph[2:], style='List Bullet')
                                elif paragraph.startswith('> *%{color:'):
                                    p = document.add_paragraph(paragraph[paragraph.find('}')+1:-2], style='Quote')
                                else:
                                    p = document.add_paragraph(paragraph, style='Normal')
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        else:
                            p = document.add_paragraph(test.procedure, style='Normal')
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                except MultiValueDictKeyError:
                    pass

                # expected
                try:
                    if request.POST['expected']:
                        document.add_heading('Ожидаемый результат', level=3)

                        if convert_textile:
                            paragraphs = test.expected.split('\r\n')
                            for paragraph in paragraphs:
                                if paragraph.startswith('#'):
                                    if paragraph.startswith('###'):
                                        p = document.add_paragraph(paragraph[4:], style='List Number 3')
                                    elif paragraph.startswith('##'):
                                        p = document.add_paragraph(paragraph[3:], style='List Number 2')
                                    else:
                                        p = document.add_paragraph(paragraph[2:], style='List Number')
                                elif paragraph.startswith('*'):
                                    if paragraph.startswith('***'):
                                        p = document.add_paragraph(paragraph[4:], style='List Bullet 3')
                                    elif paragraph.startswith('**'):
                                        p = document.add_paragraph(paragraph[3:], style='List Bullet 2')
                                    else:
                                        p = document.add_paragraph(paragraph[2:], style='List Bullet')
                                elif paragraph.startswith('> *%{color:'):
                                    p = document.add_paragraph(paragraph[paragraph.find('}')+1:-2], style='Quote')
                                else:
                                    p = document.add_paragraph(paragraph, style='Normal')
                                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                        else:
                            p = document.add_paragraph(test.expected, style='Normal')
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                except MultiValueDictKeyError:
                    pass

                # images
                try:
                    if request.POST['images']:
                        images = TestImage.objects.filter(test=test).order_by('id')
                        if images:
                            document.add_heading('Изображения', level=3)
                            for image in images:
                                document.add_paragraph(image.name, style='Subtitle')
                                document.add_picture(image.image, width=Inches(5))
                except MultiValueDictKeyError:
                    pass

                # links
                try:
                    if request.POST['links']:
                        links = TestLink.objects.filter(test=test).order_by('id')
                        if links:
                            document.add_heading('Ссылки', level=3)
                            for link in links:
                                document.add_paragraph(link.name, style='Caption')
                                document.add_paragraph(link.url, style='List Bullet')
                except MultiValueDictKeyError:
                    pass

                # checklists
                try:
                    if request.POST['checklists']:
                        checklists = TestChecklist.objects.filter(test=test).order_by('id')
                        if checklists:
                            document.add_heading('Чек-листы', level=3)
                            for checklist in checklists:
                                document.add_paragraph(checklist.name, style='Caption')
                                checklist_items = TestChecklistItem.objects.filter(checklist=checklist).order_by('id')
                                for checklist_item in checklist_items:
                                    document.add_paragraph(checklist_item.name, style='List Bullet')
                except MultiValueDictKeyError:
                    pass

                # configs
                try:
                    if request.POST['configs']:
                        configs = TestConfig.objects.filter(test=test).order_by('id')
                        if configs:
                            document.add_heading('Конфигурации', level=3)
                            for config in configs:
                                document.add_paragraph(config.name, style='Caption')
                                config.config = config.config.replace('\r', '')
                                table = document.add_table(rows=1, cols=1)
                                table.style = 'Table Grid'
                                shade_cells([table.cell(0, 0)], "#e3e8ec")
                                table.cell(0, 0).text = config.config
                except MultiValueDictKeyError:
                    pass

                # comments
                try:
                    if request.POST['comments']:
                        comments = TestComment.objects.filter(test=test).order_by('id')
                        if comments:
                            document.add_heading('Комментарии', level=3)
                            for comment in comments:
                                document.add_paragraph(comment.name, style='Caption')
                                document.add_paragraph(comment.text, style='List Bullet')
                except MultiValueDictKeyError:
                    pass

        testplan_filename = settings.MEDIA_ROOT + '/testplan_' + str(testplan.id) + '.docx'
        document.save(testplan_filename)

        file_path = os.path.join(settings.MEDIA_ROOT, testplan_filename)
        if os.path.exists(file_path):
            with open(file_path, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type="application/vnd.ms-word")
                response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_path)
                return response
        raise Http404


@login_required
def build_protocol(request):
    if request.method == 'POST':
        protocol = get_object_or_404(Protocol, id=request.POST['protocol_id'])
        profile = DocxProfile.objects.get(id=request.POST['profile_id'])

        document = build_document(profile=profile)
        sections = document.sections
        for section in sections:
            section.left_margin = Cm(2)
            section.right_margin = Cm(1.5)
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)

        document.add_heading('Протокол испытаний', level=1)

        # devices information
        try:
            if request.POST['devices']:
                document.add_heading('Состав тестируемого оборудования', level=2)

                protocol_devices = ProtocolDevice.objects.filter(protocol=protocol).order_by('id')
                for protocol_device in protocol_devices:

                    table = document.add_table(rows=0, cols=2)
                    row_cells = table.add_row().cells
                    paragraph = row_cells[0].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_text('Производитель:')
                    paragraph = row_cells[1].paragraphs[0]
                    run = paragraph.add_run()
                    run.add_text(protocol_device.device.vendor.name)

        except MultiValueDictKeyError:
            pass

        # results table
        try:
            if request.POST['results_table']:

                document.add_heading('Результаты тестов', level=2)

                table = document.add_table(rows=0, cols=4)
                row_cells = table.add_row().cells
                paragraph = row_cells[0].paragraphs[0]
                paragraph.style = 'Heading 2'
                run = paragraph.add_run()
                run.add_text('№')
                paragraph = row_cells[1].paragraphs[0]
                paragraph.style = 'Heading 2'
                run = paragraph.add_run()
                run.add_text('Название теста')
                paragraph = row_cells[2].paragraphs[0]
                paragraph.style = 'Heading 2'
                run = paragraph.add_run()
                run.add_text('Результат')
                paragraph = row_cells[3].paragraphs[0]
                paragraph.style = 'Heading 2'
                run = paragraph.add_run()
                run.add_text('Комментарии')

                categories = Category.objects.filter(testplan=protocol.testplan).order_by('priority')

                for i, category in enumerate(categories):
                    row_cells = table.add_row().cells

                    paragraph = row_cells[0].paragraphs[0]
                    paragraph.style = 'Heading 2'
                    run = paragraph.add_run()
                    run.add_text(str(i+1))

                    paragraph = row_cells[1].paragraphs[0]
                    paragraph.style = 'Heading 2'
                    run = paragraph.add_run()
                    run.add_text(category.name)

                    row_cells[1].merge(row_cells[2]).merge(row_cells[3])

                    tests = Test.objects.filter(category=category).order_by('priority')
                    for j, test in enumerate(tests):
                        row_cells = table.add_row().cells

                        paragraph = row_cells[0].paragraphs[0]
                        run = paragraph.add_run()
                        run.add_text(str(i + 1) + '.' + str(j + 1))

                        paragraph = row_cells[1].paragraphs[0]
                        run = paragraph.add_run()
                        run.add_text(test.name)

                        paragraph = row_cells[2].paragraphs[0]
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        run = paragraph.add_run()

                        try:
                            obj = ProtocolTestResult.objects.get(protocol=protocol, test=test)
                            if obj.result == 1:
                                result = 'Пропущен'
                                run.add_text(result)

                            elif obj.result == 2:
                                result = 'X'
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                                run.add_text(result)

                            elif obj.result == 3:
                                result = u'\u00b1'
                                run.font.bold = True
                                run.add_text(result)

                            else:
                                result = u'\u221a'
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)
                                run.add_text(result)

                            paragraph = row_cells[3].paragraphs[0]
                            run = paragraph.add_run()
                            issues = TestResultIssue.objects.filter(result=obj).order_by('id')
                            for issue in issues:
                                run.add_text(issue.text)

                            comments = TestResultComment.objects.filter(result=obj).order_by('id')
                            for comment in comments:
                                run.add_text(comment.text)

                        except ProtocolTestResult.DoesNotExist:
                            result = 'Пропущен'
                            run.add_text(result)

        except MultiValueDictKeyError:
            pass

        # list of issues
        try:
            if request.POST['issues']:
                issues = []
                test_results = ProtocolTestResult.objects.filter(protocol=protocol).order_by('id')
                for test_result in test_results:
                    test_result_issues = TestResultIssue.objects.filter(result=test_result)
                    for test_result_issue in test_result_issues:
                        issues.append(test_result_issue.text)
                if len(issues) > 0:
                    document.add_paragraph('По результатам испытаний выявлены следующие замечания:', style='Normal')
                    for issue in issues:
                        document.add_paragraph(issue, style='List Number')
                else:
                    document.add_paragraph('По результатам испытаний замечаний не выявлено.', style='Normal')
        except MultiValueDictKeyError:
            pass

        protocol_filename = settings.MEDIA_ROOT + '/protocol/protocol_' + str(protocol.id) + '.docx'
        document.save(protocol_filename)

        file_path = os.path.join(settings.MEDIA_ROOT, protocol_filename)
        if os.path.exists(file_path):
            with open(file_path, 'rb') as fh:
                response = HttpResponse(fh.read(), content_type="application/vnd.ms-word")
                response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_path)
                return response
        raise Http404

    else:
        raise Http404
