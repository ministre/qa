from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from testplan.models import Testplan, Category, Test, TestConfig
from django.shortcuts import get_object_or_404
from django.http import HttpResponseRedirect
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT


@login_required
def make_testplan(request):
    if request.method == 'POST':
        testplan_id = request.POST['testplan_id']
        testplan = get_object_or_404(Testplan, id=testplan_id)
        categories = Category.objects.filter(testplan=testplan).order_by('id')

        document = Document()

        style = document.styles.add_style('Header Table', WD_STYLE_TYPE.TABLE)
        style.base_style = document.styles['Table Grid']
        style.font.name = 'Cambria'
        style.font.size = Pt(11)
        style.paragraph_format.space_before = Pt(2)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        section = document.sections
        section[0].left_margin = Cm(2.5)
        section[0].right_margin = Cm(1)
        section[0].top_margin = Cm(4.5)
        section[0].bottom_margin = Cm(1)

        header = document.sections[0].header
        # paragraph = header.paragraphs[0]
        table = header.add_table(rows=0, cols=3, width=Cm(19.5))
        table.style = 'Header Table'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        row_cells = table.add_row().cells
        paragraph = row_cells[0].paragraphs[0]
        run = paragraph.add_run()
        run.add_picture('b2c.png', width=Inches(1.25))
        paragraph = row_cells[1].paragraphs[0]
        run = paragraph.add_run()
        run.add_text(testplan.name)
        row_cells = table.add_row().cells
        paragraph = row_cells[0].paragraphs[0]
        run = paragraph.add_run()
        run.add_text('Редакция: ' + testplan.version)
        paragraph = row_cells[1].paragraphs[0]
        run = paragraph.add_run()
        run.add_text('МРФ "Центр" ПАО "Ростелеком"')

        a = table.cell(0, 1)
        b = table.cell(0, 2)
        a.merge(b)

        table.cell(0, 0).width = Cm(5)
        table.cell(1, 0).width = Cm(5)
        table.cell(0, 1).width = Cm(10.5)
        table.cell(1, 1).width = Cm(10.5)
        table.cell(0, 2).width = Cm(4)
        table.cell(1, 2).width = Cm(4)
        table.cell(0, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        style = document.styles['Heading 1']
        style.paragraph_format.space_before = Pt(5)
        style.font.size = Pt(16)
        style.font.color.rgb = RGBColor(0x77, 0x00, 0xff)

        style = document.styles['Heading 2']
        style.font.size = Pt(16)
        style.font.color.rgb = RGBColor(0x77, 0x00, 0xff)

        style = document.styles['TOCHeading']
        style.font.size = Pt(14)
        style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        style.paragraph_format.space_before = Pt(7)
        style.paragraph_format.space_after = Pt(7)

        style = document.styles['List Bullet']
        style.paragraph_format.space_before = Pt(2)
        style.paragraph_format.space_after = Pt(2)

        style = document.styles['Body Text']
        style.paragraph_format.space_before = Pt(2)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.alignment = 3

        style = document.styles['Body Text 3']
        style.font.name = 'Courier New'

        style = document.styles['Table Grid']
        style.font.name = 'Courier New'
        style.font.size = Pt(8)
        style.paragraph_format.space_before = Pt(9)
        style.paragraph_format.space_after = Pt(9)

        def shade_cells(cells, shade):
            for cell in cells:
                tcPr = cell._tc.get_or_add_tcPr()
                tcVAlign = OxmlElement("w:shd")
                tcVAlign.set(qn("w:fill"), shade)
                tcPr.append(tcVAlign)

        for i, category in enumerate(categories):
            document.add_heading(str(i+1) + '. ' + category.name, level=1)

            tests = Test.objects.filter(category=category).order_by('id')
            for j, test in enumerate(tests):
                document.add_heading(str(i+1) + '.' + str(j+1) + '. ' + test.name, level=2)

                '''
                                # test description
                                table = document.add_table(rows=3, cols=2)

                                table.style = 'Table Grid'

                                table.allow_autofit = False
                                for cell in table.columns[0].cells:
                                    cell.width = Cm(2.5)
                                for cell in table.columns[1].cells:
                                    cell.width = Cm(15.5)

                                table.cell(0, 0).text = 'Цель:'
                                table.cell(1, 0).text = 'Процедура:'
                                table.cell(2, 0).text = 'Ожидаемый результат:'

                                table.cell(0, 1).text = test.purpose
                                table.cell(1, 1).text = test.procedure

                                table.cell(2, 1).text = extraspace_filter(test.expected)
                '''

                # purpose
                document.add_paragraph('Цель', style='TOCHeading')
                document.add_paragraph(test.purpose, style='Body Text')

                # procedure
                document.add_paragraph('Процедура', style='TOCHeading')
                p_num = 0
                paragraphs = test.procedure.split('\r\n')
                for paragraph in paragraphs:
                    if paragraph.startswith('#'):
                        p_num += 1
                        document.add_paragraph(str(p_num) + '. ' + paragraph[2:], style='Body Text')
                    elif paragraph.startswith('**'):
                        document.add_paragraph(paragraph[3:], style='List Bullet')
                    else:
                        document.add_paragraph(paragraph)

                # expected
                document.add_paragraph('Ожидаемый результат', style='TOCHeading')
                p_num = 0
                paragraphs = test.expected.split('\r\n')
                for paragraph in paragraphs:
                    if paragraph.startswith('#'):
                        p_num += 1
                        document.add_paragraph(str(p_num) + '. ' + paragraph[2:], style='Body Text')
                    elif paragraph.startswith('**'):
                        document.add_paragraph(paragraph[3:], style='List Bullet')
                    else:
                        document.add_paragraph(paragraph)

                # config
                configs = TestConfig.objects.filter(test=test).order_by('id')
                if configs.count():
                    document.add_paragraph('Конфигурация', style='TOCHeading')

                    for config in configs:
                        config.config = config.config.replace('\r', '')

                        table = document.add_table(rows=1, cols=1)
                        table.style = 'Table Grid'
                        shade_cells([table.cell(0, 0)], "#e3e8ec")
                        table.cell(0, 0).text = config.config

                document.add_page_break()

        document.save('demo.docx')

        return HttpResponseRedirect('/testplan/' + testplan_id + '/')
    else:
        return HttpResponseRedirect('/testplan/')


def extraspace_filter(ctx):
    ctx = ctx.replace("\r\n** ", "\n** ")
    ctx = ctx.replace("\r\n* ", "\n* ")
    ctx = ctx.replace("\r\n# ", "\n# ")
    ctx = ctx.replace("\r", "")
    return ctx
